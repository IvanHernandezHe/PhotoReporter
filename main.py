import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime

def seleccionar_imagen():
    #ruta_imagen = 'files/logo.png'
    ruta_imagen = 'files/logo3.png'
    return ruta_imagen

def seleccionar_imagenes():
    root = tk.Tk()
    root.withdraw()  # Oculta la ventana principal de tkinter

    # Abre el explorador de archivos para seleccionar múltiples imágenes
    rutas_imagenes = filedialog.askopenfilenames(title="Seleccionar imágenes", filetypes=[("Imágenes", "*.jpg;*.jpeg;*.png")])

    return rutas_imagenes

def configurar_margenes(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)
        section.header_distance = Inches(0.018)

def agregar_encabezado(doc, logo_path):

    header = doc.sections[0].header

    header_table = header.add_table(rows=1, cols=2, width=Inches(12))
    header_table.style = 'Light Shading Accent 1'

    cell_logo = header_table.cell(0, 0)
    paragraph_logo = cell_logo.paragraphs[0]
    paragraph_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    #paragraph_logo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
    run_logo = paragraph_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(2.9))
    cell_logo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    cell_info = header_table.cell(0, 1)
    paragraph_info = cell_info.paragraphs[0]
    paragraph_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph_info.add_run('REPORTE FOTOGRÁFICO\n COMERCIALIZADORA DE PRODUCTOS Y SERVICIOS\n DE REFRIGERACIÓN RECAR SA DE CV')
    run.bold = True
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor(13, 85, 137)
    cell_info.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def agregar_tabla_informacion(doc, textoAsignacion, fechaHoy):
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Medium List 1 Accent 1'

    hdr_cells = info_table.rows[0].cells

    # Asignación de texto y estilo a la primera celda
    p0 = hdr_cells[0].paragraphs[0]
    run0 = p0.add_run('Asignación: ' + textoAsignacion)
    run0.font.name = 'Segoe UI'
    run0.font.size = Pt(10)
    run0.bold = False
    run0.font.color.rgb = RGBColor(13, 85, 137)

    # Asignación de texto y estilo a la segunda celda
    p1 = hdr_cells[1].paragraphs[0]
    run1 = p1.add_run('Fecha: ' + fechaHoy)
    run1.font.name = 'Segoe UI'
    run1.font.size = Pt(10)
    run1.font.color.rgb = RGBColor(13, 85, 137)
    run1.bold = False

    # Ajustar el ancho de las celdas (opcional)
    hdr_cells[0].width = Inches(8)  # Ajusta el valor según sea necesario
    hdr_cells[1].width = Inches(1.6)  # Ajusta el valor según sea necesario

    # Alinear el texto 
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Centrar verticalmente el texto en las celdas
    hdr_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    hdr_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

from docx.shared import Inches

def agregar_tabla_con_imagenes(doc, rutas_imagenes):
    num_imagenes = len(rutas_imagenes)
    
    # Determinar el tamaño de la tabla
    if num_imagenes <= 2:
        rows = 2
        cols = 2
    elif num_imagenes <= 4:
        rows = 3
        cols = 2
    else:
        rows = ((num_imagenes - 1) // 3) + 2  # +2 porque dejamos la primera fila vacía y contamos la fila adicional necesaria
        cols = 3
    
    #4*4 y 6*5

    # Crear tabla
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light List Accent 1'

    # Calcular el ancho y alto de las celdas
    if num_imagenes <= 4:
        cell_width = Inches(3)
        cell_height = Inches(4)
    else:
        cell_width = Inches(1.5)
        cell_height = Inches(2.6)

    index = 0
    # Comenzar desde la segunda fila (índice 1)
    for i, row in enumerate(table.rows):
        if i == 0:
            continue  # Saltar la primera fila
        for cell in row.cells:
            if index < len(rutas_imagenes):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(rutas_imagenes[index], width=cell_width, height=cell_height)
                index += 1


def guardar_documento(doc, filename):
    doc.save(filename)
    print(f"Documento guardado como '{filename}'")

def convertir_a_pdf(input_file, output_file):
    # pypandoc.convert_file(input_file, 'pdf', outputfile=output_file)
    print(f"Documento guardado como '{output_file}'")



def main():

    #Se obtiene una imagen para logo
    logo_path = seleccionar_imagen()

    #se obtienen las imagenes a agregar en el reporte
    rutas_imagenes = seleccionar_imagenes()

    #Se obtiene el titulo del reporte
    textoAsignacion = 'Instalación de purga en red de aire general de linea D de la nave 1'

    #Se obtiene la fecha del día actual
    current_dateTime = datetime.now()
    fechaHoy = f"{current_dateTime.day:02d}/{current_dateTime.month:02d}/{current_dateTime.year}"

    #Se crea un documento word
    doc = Document()
    configurar_margenes(doc)

    #Se crea header
    agregar_encabezado(doc, logo_path)
    
    #Se crea infomación del reporte con titulo y día
    agregar_tabla_informacion(doc,textoAsignacion,fechaHoy)


    #Se crea tabla con imagenes
    doc.add_paragraph()
    agregar_tabla_con_imagenes(doc, rutas_imagenes[:])  # Excluimos el primer elemento (logo) de las imágenes



    docx_filename = 'ReporteConEncabezadoYMargenes.docx'
    pdf_filename = 'ReporteConEncabezadoYMargenes.pdf'

    guardar_documento(doc, docx_filename)
    convertir_a_pdf(docx_filename, pdf_filename)

if __name__ == "__main__":
    main()
